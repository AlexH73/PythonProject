import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from terms import terms  # Импортируем словарь terms из файла terms.py

def export_to_excel(terms, filename="terms.xlsx"):
    """Экспортирует данные из словаря terms в файл Excel с улучшенным форматированием."""
    workbook = openpyxl.Workbook()
    sheet = workbook.active

    # Настройка стилей
    header_font = Font(bold=True, color="FFFFFF")  # Жирный белый текст
    header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")  # Синий фон
    alignment = Alignment(horizontal="center", vertical="center")
    title_font = Font(bold=True, size=14)  # Шрифт для заголовка

    row_index = 1  # Начальная строка

    for category, term_list in terms.items():
        # Добавляем пустую строку для отступа перед заголовком
        if row_index > 1:  # Чтобы не добавлять пустую строку перед первым заголовком
            row_index += 1
        # Добавляем заголовок категории
        sheet.merge_cells(start_row=row_index, start_column=1, end_row=row_index, end_column=3)
        title_cell = sheet.cell(row=row_index, column=1)
        title_cell.value = category
        title_cell.font = title_font
        title_cell.alignment = alignment
        row_index += 1  # Переходим к следующей строке

        # Добавляем заголовки таблицы
        headers = ["Русский", "Немецкий", "Английский"]
        for col_index, header in enumerate(headers, start=1):
            cell = sheet.cell(row=row_index, column=col_index)
            cell.value = header
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = alignment

        row_index += 1  # Переходим к строкам с данными

        # Заполняем строки таблицы
        for term in term_list:
            sheet.cell(row=row_index, column=1, value=term.get('Русский', 'N/A'))
            sheet.cell(row=row_index, column=2, value=term.get('Немецкий', 'N/A'))
            sheet.cell(row=row_index, column=3, value=term.get('Английский', 'N/A'))

            # Выравниваем данные
            for col_index in range(1, 4):
                cell = sheet.cell(row=row_index, column=col_index)
                cell.alignment = alignment

            row_index += 1  # Следующая строка

        # Добавляем отступ между категориями
        row_index += 1

    # Настройка ширины колонок
    for column_cells in sheet.columns:
        max_length = 0
        # Получаем букву колонки через объект столбца
        column_letter = column_cells[0].coordinate[:1]

        for cell in column_cells:
            # Пропускаем объединенные ячейки
            if isinstance(cell, openpyxl.cell.cell.MergedCell):
                continue
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except Exception:
                pass

        adjusted_width = max_length + 2  # Добавляем небольшой отступ
        sheet.column_dimensions[column_letter].width = adjusted_width

    # Границы для всех ячеек
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )
    for row in sheet.iter_rows():
        for cell in row:
            cell.border = thin_border

    # Сохраняем файл
    workbook.save(filename)
    print(f"Данные успешно экспортированы в файл '{filename}'")


def export_to_word(terms, filename="terms.docx"):
    """Экспортирует данные из словаря terms в файл Word."""
    document = Document()

    for category, term_list in terms.items():
        # Добавляем заголовок раздела
        document.add_heading(category, level=1)

        # Создаем таблицу с фиксированными колонками
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.allow_autofit = False
        table.columns[0].width = Inches(2)
        table.columns[1].width = Inches(2)
        table.columns[2].width = Inches(2)

        # Заполняем заголовок таблицы
        header_cells = table.rows[0].cells
        header_cells[0].text = "Русский"
        header_cells[1].text = "Немецкий"
        header_cells[2].text = "Английский"
        for cell in header_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Заполняем строки таблицы
        for term in term_list:
            row = table.add_row().cells
            row[0].text = term.get('Русский', 'N/A')
            row[1].text = term.get('Немецкий', 'N/A')
            row[2].text = term.get('Английский', 'N/A')
            for cell in row:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Добавляем разделитель между таблицами
        document.add_paragraph("\n")

    document.save(filename)
    print(f"Данные успешно экспортированы в файл '{filename}'")

if __name__ == "__main__":
    export_to_excel(terms)
    export_to_word(terms)